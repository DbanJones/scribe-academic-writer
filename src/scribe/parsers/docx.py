"""DOCX text extraction using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def extract(path: Path) -> str:
    """Extract text from a DOCX file as markdown-ish text."""
    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            parts.append(f"# {text}")
        elif "heading 2" in style_name:
            parts.append(f"## {text}")
        elif "heading 3" in style_name:
            parts.append(f"### {text}")
        elif "list" in style_name:
            parts.append(f"- {text}")
        else:
            parts.append(text)

    # Extract tables as markdown
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if rows:
            # Header row
            parts.append("| " + " | ".join(rows[0]) + " |")
            parts.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for row in rows[1:]:
                parts.append("| " + " | ".join(row) + " |")
            parts.append("")

    return "\n\n".join(parts)
