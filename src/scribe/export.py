"""Markdown-to-Word (.docx) exporter.

Produces a nicely formatted .docx from the markdown that Scribe's pipelines
emit. Handles:

- Heading hierarchy (# -> H1, ## -> H2, etc.)
- Paragraphs and blank-line separation
- Bulleted and numbered lists (nested up to 4 levels)
- Markdown tables
- Inline formatting: **bold**, *italic*, `code`, [link](url)
- Figure captions (lines starting "Figure N:" or "Table N:")
- Visual suggestion placeholders (![SUGGEST: ...](suggest))
- Block quotes (> ...)
- Horizontal rules (---)
- Bibliography section (detects "# References" / "# Bibliography")

Strategy: a small line-by-line parser, not a general markdown engine. The
output is produced by Scribe's pipelines so we can rely on clean structure.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

logger = logging.getLogger(__name__)


# --- Regex patterns for inline formatting ---

# Order matters: handle more specific patterns before the generic ones.
# Each pattern captures a single stylistic run.
_INLINE_TOKEN = re.compile(
    r"""
    (?P<bold>\*\*(?P<bold_text>.+?)\*\*)            # **bold**
    | (?P<italic>(?<![\\*])\*(?!\*)(?P<italic_text>[^*\n]+?)\*(?!\*))  # *italic*
    | (?P<code>`(?P<code_text>[^`\n]+?)`)           # `code`
    | (?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))  # [text](url)
    """,
    re.VERBOSE,
)

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_BULLET_RE = re.compile(r"^(?P<indent>\s*)[-*+]\s+(?P<text>.+)$")
_NUM_LIST_RE = re.compile(r"^(?P<indent>\s*)(?P<num>\d+)[.)]\s+(?P<text>.+)$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
_TABLE_DIV_RE = re.compile(r"^\|[\s\-:|]+\|\s*$")
_HR_RE = re.compile(r"^(\s*[-*_]\s*){3,}\s*$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_FIGURE_CAPTION_RE = re.compile(r"^(Figure|Table)\s+\d+[:.]\s+", re.IGNORECASE)
_SUGGEST_RE = re.compile(r"!\[SUGGEST:\s*(?P<desc>.+?)\]\(suggest\)")
_BIB_ENTRY_RE = re.compile(r"^[A-Z][A-Za-z\-',. ]+\s*\(\d{4}[a-z]?\)")


# --- Styling ---


@dataclass
class ExportStyle:
    """Typographic settings for the exported document."""
    body_font: str = "Calibri"
    heading_font: str = "Calibri"
    serif_font: str = "Cambria"   # optional override for a more traditional look
    mono_font: str = "Consolas"
    body_size_pt: float = 11.0
    line_spacing: float = 1.3
    heading_colour: tuple = (0x1F, 0x2D, 0x5B)  # deep ink-blue
    link_colour: tuple = (0x8A, 0x9E, 0xFF)
    code_background: tuple = (0xF2, 0xF2, 0xF6)


_IN_BIBLIOGRAPHY: bool = False  # module-level toggle, reset per export


def markdown_to_docx(
    markdown_text: str,
    output_path: Path,
    *,
    title: str | None = None,
    style: ExportStyle | None = None,
) -> Path:
    """Convert a markdown string to a .docx file.

    Args:
        markdown_text: The markdown content.
        output_path: Where to write the .docx.
        title: Optional document title. If omitted, the first H1 is used.
        style: ExportStyle overrides.

    Returns:
        The path to the written file.
    """
    style = style or ExportStyle()
    doc = Document()
    _configure_document_styles(doc, style)

    # Set sensible page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # If the caller gave us a title, add it up front as a Title-styled run.
    # Otherwise the first H1 in the body acts as the title.
    if title:
        title_para = doc.add_paragraph()
        title_para.style = doc.styles["Title"]
        run = title_para.add_run(title)
        run.font.name = style.heading_font
        run.font.size = Pt(22)

    _render_lines(doc, markdown_text.splitlines(), style)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Wrote %s", output_path)
    return output_path


# --- Core parser ---


def _render_lines(doc: Document, lines: list[str], style: ExportStyle) -> None:
    global _IN_BIBLIOGRAPHY
    _IN_BIBLIOGRAPHY = False

    i = 0
    n = len(lines)
    saw_first_h1 = False

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if _HR_RE.match(line):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        h_match = _HEADING_RE.match(line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            # Track whether we've moved into a bibliography section,
            # so the subsequent paragraphs can be rendered as hanging indents.
            lowered = text.lower()
            if level <= 2 and (
                "references" in lowered or "bibliography" in lowered or "works cited" in lowered
            ):
                _IN_BIBLIOGRAPHY = True
            elif level <= 2:
                _IN_BIBLIOGRAPHY = False

            # Promote the first H1 to Title style if we didn't have an explicit title.
            if level == 1 and not saw_first_h1 and not _has_title_paragraph(doc):
                _add_title(doc, text, style)
                saw_first_h1 = True
            else:
                _add_heading(doc, text, level, style)
            i += 1
            continue

        # Table (detect by looking ahead at the divider row; tolerate blank
        # lines between rows, since some docx extractors separate every
        # paragraph with blank lines).
        if _TABLE_ROW_RE.match(line):
            j = i + 1
            while j < n and not lines[j].strip():
                j += 1
            if j < n and _TABLE_DIV_RE.match(lines[j].rstrip()):
                consumed = _add_table(doc, lines, i, style)
                i += consumed
                continue

        # Block quote (consume consecutive > lines)
        if _BLOCKQUOTE_RE.match(line):
            consumed = _add_blockquote(doc, lines, i, style)
            i += consumed
            continue

        # Bulleted list (consume consecutive bullets)
        if _BULLET_RE.match(line):
            consumed = _add_list(doc, lines, i, style, numbered=False)
            i += consumed
            continue

        # Numbered list
        if _NUM_LIST_RE.match(line):
            consumed = _add_list(doc, lines, i, style, numbered=True)
            i += consumed
            continue

        # Visual suggestion
        if _SUGGEST_RE.search(line):
            _add_figure_suggestion(doc, line, style)
            i += 1
            continue

        # Figure / table caption
        if _FIGURE_CAPTION_RE.match(line):
            _add_figure_caption(doc, line, style)
            i += 1
            continue

        # Bibliography entry (in a References section)
        if _IN_BIBLIOGRAPHY and _BIB_ENTRY_RE.match(line):
            _add_bib_entry(doc, line, style)
            i += 1
            continue

        # Default: paragraph (may span multiple non-blank lines joined by \n)
        # A paragraph ends at a blank line, heading, list, table, blockquote, etc.
        buf = [line]
        i += 1
        while i < n:
            next_line = lines[i].rstrip()
            if not next_line.strip():
                break
            if _HEADING_RE.match(next_line):
                break
            if _BULLET_RE.match(next_line) or _NUM_LIST_RE.match(next_line):
                break
            if _TABLE_ROW_RE.match(next_line):
                break
            if _BLOCKQUOTE_RE.match(next_line):
                break
            if _HR_RE.match(next_line):
                break
            buf.append(next_line)
            i += 1

        _add_paragraph(doc, " ".join(buf), style)


# --- Block renderers ---


def _add_heading(doc: Document, text: str, level: int, style: ExportStyle) -> None:
    # python-docx has Heading 1..9
    para = doc.add_heading(level=min(level, 9))
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    _add_inline_runs(para, text, style, base_font=style.heading_font,
                     size_pt=_heading_size(level), colour=style.heading_colour,
                     bold=True)


def _heading_size(level: int) -> float:
    return {1: 20.0, 2: 16.0, 3: 13.5, 4: 12.0, 5: 11.5, 6: 11.0}.get(level, 11.0)


def _add_title(doc: Document, text: str, style: ExportStyle) -> None:
    para = doc.add_paragraph()
    para.style = doc.styles["Title"]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(18)
    run = para.add_run(text)
    run.font.name = style.heading_font
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*style.heading_colour)


def _has_title_paragraph(doc: Document) -> bool:
    for para in doc.paragraphs:
        if para.style.name == "Title" and para.text.strip():
            return True
    return False


def _add_paragraph(doc: Document, text: str, style: ExportStyle) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = style.line_spacing
    _add_inline_runs(para, text, style)


def _add_figure_caption(doc: Document, text: str, style: ExportStyle) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(text)
    run.font.name = style.body_font
    run.font.size = Pt(style.body_size_pt - 1)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)


def _add_figure_suggestion(doc: Document, line: str, style: ExportStyle) -> None:
    """Render ![SUGGEST: ...](suggest) as a boxed placeholder."""
    match = _SUGGEST_RE.search(line)
    desc = match.group("desc") if match else line
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    # Put a top+bottom border on the paragraph for a boxed effect
    p_pr = para._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "8A9EFF")
        borders.append(b)
    p_pr.append(borders)

    # Shade the paragraph background lightly
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F5F7FF")
    p_pr.append(shd)

    label = para.add_run("[suggested visual]  ")
    label.font.name = style.body_font
    label.font.size = Pt(style.body_size_pt - 1)
    label.font.bold = True
    label.font.color.rgb = RGBColor(0x5A, 0x6C, 0xC8)

    desc_run = para.add_run(desc)
    desc_run.font.name = style.body_font
    desc_run.font.size = Pt(style.body_size_pt - 1)
    desc_run.font.italic = True


def _add_horizontal_rule(doc: Document) -> None:
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    bdr.append(bottom)
    p_pr.append(bdr)


def _add_list(
    doc: Document,
    lines: list[str],
    start: int,
    style: ExportStyle,
    *,
    numbered: bool,
) -> int:
    """Consume consecutive list lines; return how many lines we used."""
    i = start
    n = len(lines)
    consumed = 0

    while i < n:
        raw = lines[i].rstrip()
        if not raw.strip():
            break

        if numbered:
            m = _NUM_LIST_RE.match(raw)
            if not m and not _BULLET_RE.match(raw):
                break
        else:
            m = _BULLET_RE.match(raw)
            if not m and not _NUM_LIST_RE.match(raw):
                break

        # Accept mixed bullet / numbered mid-list: follow whichever matched
        m = _NUM_LIST_RE.match(raw) or _BULLET_RE.match(raw)
        indent = len(m.group("indent"))
        level = min(indent // 2, 3)
        text = m.group("text")

        style_name = "List Number" if numbered and _NUM_LIST_RE.match(raw) else "List Bullet"
        if level > 0:
            style_name = f"{style_name} {level + 1}"

        try:
            para = doc.add_paragraph(style=style_name)
        except KeyError:
            para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        _add_inline_runs(para, text, style)
        i += 1
        consumed += 1

    return consumed or 1


def _add_blockquote(
    doc: Document, lines: list[str], start: int, style: ExportStyle,
) -> int:
    i = start
    n = len(lines)
    buf: list[str] = []
    consumed = 0
    while i < n:
        raw = lines[i].rstrip()
        m = _BLOCKQUOTE_RE.match(raw)
        if not m:
            break
        buf.append(m.group(1))
        i += 1
        consumed += 1

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Left border
    p_pr = para._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "8A9EFF")
    bdr.append(left)
    p_pr.append(bdr)
    _add_inline_runs(para, " ".join(buf).strip(), style, italic=True)
    return consumed or 1


def _add_table(
    doc: Document, lines: list[str], start: int, style: ExportStyle,
) -> int:
    """Parse a markdown table starting at `start`. Returns rows consumed.

    Tolerates blank lines between table rows (common when the input was
    produced by a paragraph-oriented docx extractor).
    """
    i = start
    n = len(lines)

    def _advance_past_blanks() -> None:
        """Move i forward past any blank lines."""
        nonlocal i
        while i < n and not lines[i].strip():
            i += 1

    # Collect header row
    header_match = _TABLE_ROW_RE.match(lines[i].rstrip())
    if not header_match:
        return 1
    header_cells = _split_table_row(lines[i])
    i += 1

    # Divider row (possibly preceded by blanks)
    _advance_past_blanks()
    if i < n and _TABLE_DIV_RE.match(lines[i].rstrip()):
        i += 1

    # Body rows
    body_rows: list[list[str]] = []
    while i < n:
        _advance_past_blanks()
        if i >= n:
            break
        raw = lines[i].rstrip()
        if not _TABLE_ROW_RE.match(raw):
            break
        if _TABLE_DIV_RE.match(raw):
            i += 1
            continue

        # Peek ahead: if the next non-blank line is a divider, this is the
        # header of a NEW table, not another row of this one.
        peek = i + 1
        while peek < n and not lines[peek].strip():
            peek += 1
        if peek < n and _TABLE_DIV_RE.match(lines[peek].rstrip()):
            break

        body_rows.append(_split_table_row(raw))
        i += 1

    cols = len(header_cells)
    tbl = doc.add_table(rows=1 + len(body_rows), cols=cols)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for c, cell_text in enumerate(header_cells):
        cell = tbl.rows[0].cells[c]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        run = para.add_run(cell_text.strip())
        run.bold = True
        run.font.name = style.body_font
        run.font.size = Pt(style.body_size_pt - 0.5)

    # Body
    for r, row in enumerate(body_rows, start=1):
        for c in range(cols):
            cell_text = row[c] if c < len(row) else ""
            cell = tbl.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            _add_inline_runs(para, cell_text.strip(), style,
                             size_override=style.body_size_pt - 0.5)

    # Trailing paragraph spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)

    return (i - start) or 1


def _split_table_row(raw: str) -> list[str]:
    """Split a '| a | b | c |' row into ['a', 'b', 'c'] (preserving markdown)."""
    # Strip leading/trailing pipes, then split on |, honouring escaped \|.
    stripped = raw.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    # Simple split is enough here; escaped pipes are rare in our output.
    return [c.strip() for c in stripped.split("|")]


def _add_bib_entry(doc: Document, text: str, style: ExportStyle) -> None:
    """Render a bibliography entry with hanging indent."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.first_line_indent = Cm(-1.0)  # hanging
    para.paragraph_format.space_after = Pt(6)
    _add_inline_runs(para, text, style)


# --- Inline runs ---


def _add_inline_runs(
    para,
    text: str,
    style: ExportStyle,
    *,
    base_font: str | None = None,
    size_pt: float | None = None,
    colour: tuple | None = None,
    bold: bool = False,
    italic: bool = False,
    size_override: float | None = None,
) -> None:
    """Parse inline markdown and add styled runs to a paragraph."""
    if not text:
        return

    default_font = base_font or style.body_font
    default_size = size_pt or size_override or style.body_size_pt

    pos = 0
    for match in _INLINE_TOKEN.finditer(text):
        start, end = match.span()
        if start > pos:
            _add_plain_run(
                para, text[pos:start], default_font, default_size, colour,
                bold=bold, italic=italic,
            )

        if match.group("bold"):
            inner = match.group("bold_text")
            _add_plain_run(
                para, inner, default_font, default_size, colour,
                bold=True, italic=italic,
            )
        elif match.group("italic"):
            inner = match.group("italic_text")
            _add_plain_run(
                para, inner, default_font, default_size, colour,
                bold=bold, italic=True,
            )
        elif match.group("code"):
            inner = match.group("code_text")
            _add_plain_run(
                para, inner, style.mono_font, default_size,
                colour=(0x2B, 0x2B, 0x2B), bold=bold, italic=italic,
                shade=style.code_background,
            )
        elif match.group("link"):
            _add_hyperlink(para, match.group("link_text"),
                           match.group("link_url"), style, default_size)

        pos = end

    if pos < len(text):
        _add_plain_run(
            para, text[pos:], default_font, default_size, colour,
            bold=bold, italic=italic,
        )


def _add_plain_run(
    para,
    text: str,
    font_name: str,
    size_pt: float,
    colour: tuple | None,
    *,
    bold: bool = False,
    italic: bool = False,
    shade: tuple | None = None,
) -> None:
    if not text:
        return
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = RGBColor(*colour)
    # Force the East Asian font so the choice takes on all code pages
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    if shade:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), _rgb_to_hex(shade))
        rpr.append(shd)


def _add_hyperlink(
    para, text: str, url: str, style: ExportStyle, size_pt: float,
) -> None:
    """Add a clickable hyperlink run to the paragraph."""
    part = para.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # Colour and underline
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _rgb_to_hex(style.link_colour))
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    rpr.append(sz)
    run.append(rpr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    text_elem.set(qn("xml:space"), "preserve")
    run.append(text_elem)
    hyperlink.append(run)
    para._p.append(hyperlink)


def _rgb_to_hex(rgb: tuple) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# --- Document-wide styling ---


def _configure_document_styles(doc: Document, style: ExportStyle) -> None:
    """Apply body font and heading colours to the built-in styles."""
    # Body paragraph defaults
    normal = doc.styles["Normal"]
    normal.font.name = style.body_font
    normal.font.size = Pt(style.body_size_pt)
    normal.paragraph_format.line_spacing = style.line_spacing
    normal.paragraph_format.space_after = Pt(6)
    # Also set East Asian font via rFonts so it survives on non-Latin scripts
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), style.body_font)
    rfonts.set(qn("w:hAnsi"), style.body_font)
    rfonts.set(qn("w:cs"), style.body_font)

    # Heading styles
    for level in range(1, 7):
        try:
            h = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        h.font.name = style.heading_font
        h.font.size = Pt(_heading_size(level))
        h.font.color.rgb = RGBColor(*style.heading_colour)
        h.font.bold = True
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.space_before = Pt(14 - level)
        h.paragraph_format.space_after = Pt(4)


# --- Convenience entry points ---


def export_file(md_path: Path, docx_path: Path | None = None,
                title: str | None = None) -> Path:
    """Export a markdown file to a .docx sibling (or explicit target)."""
    md_path = Path(md_path)
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown not found: {md_path}")
    if docx_path is None:
        docx_path = md_path.with_suffix(".docx")
    text = md_path.read_text(encoding="utf-8")
    return markdown_to_docx(text, docx_path, title=title)


def try_export_sibling(md_path: Path) -> Path | None:
    """Export alongside the given .md file, swallowing any error.

    Used by the pipeline finishers where a .docx export is nice-to-have
    but should never abort a successful run.
    """
    try:
        return export_file(md_path)
    except Exception as e:  # noqa: BLE001 -- pipeline must not fail over this
        logger.warning("DOCX export failed for %s: %s", md_path, e)
        return None
